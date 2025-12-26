from sqlalchemy.orm import Session
from typing import Optional, List
from datetime import date
import pandas as pd
from fastapi.responses import StreamingResponse
from sqlalchemy import text
from src.schema.reports import FullHospitalData


def convert_boolean_to_yes_no(value):
    if value is None:
        return None
    if isinstance(value, str):
        if value.lower() in ['0', 'false', 'no']:
            return 'No'
        elif value.lower() in ['1', 'true', 'yes']:
            return 'Yes'
        return value
    if isinstance(value, (int, bool)):
        return 'Yes' if value else 'No'
    return value


def get_hospital_data_from_proc(
        db: Session,
        hospital_id: Optional[int] = None,
        state: Optional[str] = None,
        city: Optional[str] = None,
        district: Optional[str] = None,
        taluka: Optional[str] = None,
        from_date: Optional[date] = None,
        to_date: Optional[date] = None,
) -> List[FullHospitalData]:
    result = db.execute(
        text("""
            CALL hospital_audit.GetHospitalCompleteData1(
                :hospital_id, :state, :city, :district, :taluka, :from_date, :to_date
            )
        """),
        {
            "hospital_id": hospital_id,
            "state": state,
            "city": city,
            "district": district,
            "taluka": taluka,
            "from_date": from_date,
            "to_date": to_date,
        }
    )
    rows = result.fetchall()
    columns = result.keys()
    hospitals = []
    boolean_fields = [
        'mtp_registration', 'pndt_registration', 'organ_transplant_registration',
        'state_clinical_license', 'nabh_accreditation', 'nabl_accreditation',
        'jci_accreditation', 'iso_certification', 'bio_medical_waste_mgmt_registration',
        'fire_safety_registration', 'document_consistency_verified', 'edited_certificates_found',
        'in_house_lab', 'in_house_pharmacy', 'ambulance_service', 'reception',
        'tpa_insurance_department', 'patient_counselling_area', 'physiotherapy_department',
        'admin_office_present', 'central_oxygen_supply', 'infra_equipment_validation',
        'infrastructure_verified', 'stock_register', 'purchase_bills', 'sales_register',
        'ipd_register', 'icu_register', 'opd_register', 'operation_theater_register',
        'lab_register', 'mlc_register', 'infection_control_protocols',
        'consultant_visits_register', 'cancelled_cheque_copy', 'bagic_mou_copy',
        'new_setup_claim_verified', 'human_presence_verified', 'repeat_tpa_fraud_suspected',
        'xray_machine', 'c_arm_machine', 'ct_scan', 'mri_machine', 'usg_machine'
    ]
    for row in rows:
        row_dict = dict(zip(columns, row))
        for field in boolean_fields:
            if field in row_dict:
                row_dict[field] = convert_boolean_to_yes_no(row_dict[field])

        hospitals.append(FullHospitalData(
            basic_info={
                "hospital_id": row_dict.get("hospital_id"),
                "hospital_name": row_dict.get("hospital_name"),
                "address": row_dict.get("address"),
                "location_city": row_dict.get("location_city"),
                "state": row_dict.get("state"),
                "area_type": row_dict.get("area_type"),
                "hospital_ownership_type": row_dict.get("hospital_ownership_type"),
                "run_by": row_dict.get("run_by"),
                "hospital_type": row_dict.get("hospital_type"),
                "specialty": row_dict.get("specialty"),
                "specialty_other": row_dict.get("specialty_other"),
                "audit_status": row_dict.get("audit_status"),

                "hospital_registration_number": row_dict.get("hospital_registration_number"),
                "registered_with": row_dict.get("registered_with"),
                "registration_valid_up_to": row_dict.get("registration_valid_up_to"),
                "registration_number_of_beds": row_dict.get("registration_number_of_beds"),
                "mtp_registration": row_dict.get("mtp_registration"),
                "pndt_registration": row_dict.get("pndt_registration"),
                "organ_transplant_registration": row_dict.get("organ_transplant_registration"),
                "state_clinical_license": row_dict.get("state_clinical_license"),
                "nabh_accreditation": row_dict.get("nabh_accreditation"),
                "nabl_accreditation": row_dict.get("nabl_accreditation"),
                "jci_accreditation": row_dict.get("jci_accreditation"),
                "iso_certification": row_dict.get("iso_certification"),
                "bio_medical_waste_mgmt_registration": row_dict.get("bio_medical_waste_mgmt_registration"),
                "fire_safety_registration": row_dict.get("fire_safety_registration"),
                "other_certifications": row_dict.get("other_certifications"),
                "gst_number": row_dict.get("gst_number"),
                "tan_number": row_dict.get("tan_number"),
                "pan_number": row_dict.get("pan_number"),
                "pan_name": row_dict.get("pan_name"),
                "establishment_year": row_dict.get("establishment_year"),
                "certificate_of_incorporation": row_dict.get("certificate_of_incorporation"),
                "document_consistency_verified": row_dict.get("document_consistency_verified"),
                "edited_certificates_found": row_dict.get("edited_certificates_found"),

                "govt_schemes_run": row_dict.get("govt_schemes_run"),
                "govt_schemes_run_other": row_dict.get("govt_schemes_run_other"),
                "tpas_empaneled": row_dict.get("tpas_empaneled"),
                "tpas_empaneled_other": row_dict.get("tpas_empaneled_other"),
                "insurance_companies_empaneled": row_dict.get("insurance_companies_empaneled"),
                "corporates_empaneled": row_dict.get("corporates_empaneled"),

                "medical_superintendent_name": row_dict.get("medical_superintendent_name"),
                "medical_superintendent_working_type": row_dict.get("medical_superintendent_working_type"),
                "medical_superintendent_contact": row_dict.get("medical_superintendent_contact"),

                "pathologist_name": row_dict.get("pathologist_name"),
                "pathologist_working_type": row_dict.get("pathologist_working_type"),
                "pathologist_contact": row_dict.get("pathologist_contact"),
                "lab_technician_name": row_dict.get("lab_technician_name"),
                "lab_technician_working_type": row_dict.get("lab_technician_working_type"),
                "lab_technician_contact": row_dict.get("lab_technician_contact"),
                "radiologist_name": row_dict.get("radiologist_name"),
                "radiologist_working_type": row_dict.get("radiologist_working_type"),
                "radiologist_contact": row_dict.get("radiologist_contact"),

                "nurse_sample_registrations": row_dict.get("nurse_sample_registrations"),
                "number_of_specialist_doctors": row_dict.get("number_of_specialist_doctors"),
                "number_of_super_specialist_doctors": row_dict.get("number_of_super_specialist_doctors"),
                "number_of_mbbs_and_above_doctors": row_dict.get("number_of_mbbs_and_above_doctors"),
                "number_of_trained_nurses": row_dict.get("number_of_trained_nurses"),
                "number_of_housemen": row_dict.get("number_of_housemen"),
                "number_of_resident_doctors": row_dict.get("number_of_resident_doctors"),
                "number_of_untrained_nurses": row_dict.get("number_of_untrained_nurses"),
                "number_of_admin_staff": row_dict.get("number_of_admin_staff"),
                "number_of_other_staff": row_dict.get("number_of_other_staff"),
                "total_number_of_staff": row_dict.get("total_number_of_staff"),
                "nurse_to_patient_ratio": row_dict.get("nurse_to_patient_ratio"),

                "type_of_building": row_dict.get("type_of_building"),
                "area_in_square_feet": row_dict.get("area_in_square_feet"),
                "number_of_floors": row_dict.get("number_of_floors"),
                "number_of_beds": row_dict.get("number_of_beds"),
                "licensed_bed_count": row_dict.get("licensed_bed_count"),
                "in_house_lab": row_dict.get("in_house_lab"),
                "in_house_pharmacy": row_dict.get("in_house_pharmacy"),
                "ambulance_service": row_dict.get("ambulance_service"),
                "reception": row_dict.get("reception"),
                "tpa_insurance_department": row_dict.get("tpa_insurance_department"),
                "patient_counselling_area": row_dict.get("patient_counselling_area"),
                "physiotherapy_department": row_dict.get("physiotherapy_department"),
                "admin_office_present": row_dict.get("admin_office_present"),
                "number_of_opds": row_dict.get("number_of_opds"),
                "number_of_icus": row_dict.get("number_of_icus"),
                "total_icu_beds": row_dict.get("total_icu_beds"),
                "special_rooms": row_dict.get("special_rooms"),
                "semi_special_rooms": row_dict.get("semi_special_rooms"),
                "number_of_suites": row_dict.get("number_of_suites"),
                "number_of_general_wards": row_dict.get("number_of_general_wards"),
                "number_of_general_ward_beds": row_dict.get("number_of_general_ward_beds"),
                "types_of_operation_theaters": row_dict.get("types_of_operation_theaters"),
                "number_of_operation_theaters": row_dict.get("number_of_operation_theaters"),
                "central_oxygen_supply": row_dict.get("central_oxygen_supply"),
                "number_of_ventilators": row_dict.get("number_of_ventilators"),
                "infra_equipment_validation": row_dict.get("infra_equipment_validation"),
                "infrastructure_verified": row_dict.get("infrastructure_verified"),
                "other_infrastructure": row_dict.get("other_infrastructure"),

                "lab_name": row_dict.get("lab_name"),
                "lab_pathologist_name": row_dict.get("lab_pathologist_name"),
                "lab_pathologist_type": row_dict.get("lab_pathologist_type"),
                "qualification_of_lab_pathologist": row_dict.get("qualification_of_lab_pathologist"),
                "pathologist_own_lab_name_location": row_dict.get("pathologist_own_lab_name_location"),
                "lab_equipment_available": row_dict.get("lab_equipment_available"),
                "type_of_tests_done": row_dict.get("type_of_tests_done"),
                "tests_outsourced": row_dict.get("tests_outsourced"),
                "xray_machine": row_dict.get("xray_machine"),
                "c_arm_machine": row_dict.get("c_arm_machine"),
                "ct_scan": row_dict.get("ct_scan"),
                "mri_machine": row_dict.get("mri_machine"),
                "usg_machine": row_dict.get("usg_machine"),
                "other_diagnostic_equipment": row_dict.get("other_diagnostic_equipment"),

                "pharmacy_name": row_dict.get("pharmacy_name"),
                "pharmacy_drug_license_number": row_dict.get("pharmacy_drug_license_number"),
                "pharmacy_working_hours": row_dict.get("pharmacy_working_hours"),
                "pharmacist_name": row_dict.get("pharmacist_name"),
                "pharmacist_registration_number": row_dict.get("pharmacist_registration_number"),
                "stock_register": row_dict.get("stock_register"),
                "purchase_bills": row_dict.get("purchase_bills"),
                "sales_register": row_dict.get("sales_register"),
                "other_pharmacy_documentation": row_dict.get("other_pharmacy_documentation"),

                "ipd_register": row_dict.get("ipd_register"),
                "icu_register": row_dict.get("icu_register"),
                "opd_register": row_dict.get("opd_register"),
                "operation_theater_register": row_dict.get("operation_theater_register"),
                "lab_register": row_dict.get("lab_register"),
                "mlc_register": row_dict.get("mlc_register"),
                "infection_control_protocols": row_dict.get("infection_control_protocols"),
                "consultant_visits_register": row_dict.get("consultant_visits_register"),

                "bank_account_number": row_dict.get("bank_account_number"),
                "ifsc_code": row_dict.get("ifsc_code"),
                "cancelled_cheque_copy": row_dict.get("cancelled_cheque_copy"),

                "bagic_to_other_patients_ratio": row_dict.get("bagic_to_other_patients_ratio"),
                "average_stay": row_dict.get("average_stay"),
                "surgical_to_medical_packages_ratio": row_dict.get("surgical_to_medical_packages_ratio"),
                "bagic_mou_copy": row_dict.get("bagic_mou_copy"),
                "bagic_soc_type": row_dict.get("bagic_soc_type"),
                "bagic_empanelment_done_by": row_dict.get("bagic_empanelment_done_by"),
                "bagic_empanelment_experience": row_dict.get("bagic_empanelment_experience"),
                "bagic_claim_approval_experience": row_dict.get("bagic_claim_approval_experience"),
                "bagic_claim_verification_experience": row_dict.get("bagic_claim_verification_experience"),
                "bagic_payments_experience": row_dict.get("bagic_payments_experience"),
                "suggestion_for_bagic": row_dict.get("suggestion_for_bagic"),

                "date_of_visit": row_dict.get("date_of_visit"),
                "visit_done_by": row_dict.get("visit_done_by"),
                "attended_by": row_dict.get("attended_by"),
                "designation": row_dict.get("designation"),
                "information_given_by": row_dict.get("information_given_by"),
                "remark_by_auditor": row_dict.get("remark_by_auditor"),
                "remark_by_hospital_manager_owner": row_dict.get("remark_by_hospital_manager_owner"),
                "hospital_manager_owner_name": row_dict.get("hospital_manager_owner_name"),
                "new_setup_claim_verified": row_dict.get("new_setup_claim_verified"),
                "human_presence_verified": row_dict.get("human_presence_verified"),
                "local_market_intel_remarks": row_dict.get("local_market_intel_remarks"),
                "repeat_tpa_fraud_suspected": row_dict.get("repeat_tpa_fraud_suspected"),
            }
        ))
    return hospitals


def is_empty_value(value):
    if value is None:
        return True
    if isinstance(value, str):
        return value.strip() == ""
    if isinstance(value, (list, dict)):
        return len(value) == 0
    return False

def export_hospitals_to_excel(hospitals: List[FullHospitalData]):
    flat_data = []
    excluded_fields = ['hospital_id']

    for h in hospitals:
        row = {}
        if h.basic_info:
            basic_info_dict = h.basic_info.dict()
            # keep all keys, even if value is None/empty
            filtered_dict = {key: value for key, value in basic_info_dict.items() if key not in excluded_fields}
            row.update(filtered_dict)
        flat_data.append(row)

    if not flat_data:
        df = pd.DataFrame()
    else:
        df = pd.DataFrame(flat_data)

    stream = BytesIO()
    df.to_excel(stream, index=False)
    stream.seek(0)

    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=hospitals.xlsx"}
    )


#
# def export_hospitals_to_excel(hospitals: List[FullHospitalData]):
#     flat_data = []
#     excluded_fields = ['hospital_id']
#
#     for h in hospitals:
#         row = {}
#         if h.basic_info:
#             basic_info_dict = h.basic_info.dict()
#             filtered_dict = {}
#             for key, value in basic_info_dict.items():
#                 if key not in excluded_fields and not is_empty_value(value):
#                     filtered_dict[key] = value
#             row.update(filtered_dict)
#         flat_data.append(row)
#     flat_data = [row for row in flat_data if row]
#
#     if not flat_data:
#         df = pd.DataFrame()
#     else:
#         df = pd.DataFrame(flat_data)
#         df = df.dropna(axis=1, how='all')
#
#     stream = BytesIO()
#     df.to_excel(stream, index=False)
#     stream.seek(0)
#
#     return StreamingResponse(
#         stream,
#         media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
#         headers={"Content-Disposition": "attachment; filename=hospitals.xlsx"}
#     )

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from zipfile import ZipFile
from io import BytesIO
import os
from datetime import datetime
from typing import List, Optional


def create_professional_field_mapping():
    return {
        'hospital_id': 'Hospital ID',
        'hospital_name': 'Hospital Name',
        'address': 'Address',
        'location_city': 'City',
        'state': 'State',
        'area_type': 'Area Type',
        'hospital_ownership_type': 'Ownership Type',
        'run_by': 'Managed By',
        'owned_by': 'Owned By',
        'qualification_of_owner': 'Owner Qualification',
        'hospital_type': 'Hospital Type',
        'specialty_type': 'Specialty Type',
        'specialty': 'Specialty',
        'specialty_other': 'Other Specialty',
        'rohini_id': 'ROHINI ID',
        'unique_code': 'Unique Code',
        'audit_status': 'Audit Status',
        'hospital_registration_number': 'Hospital Registration Number',
        'registered_with': 'Registered With',
        'registration_valid_up_to': 'Registration Valid Upto',
        'registration_number_of_beds': 'Number of Beds Registered',
        'mtp_registration': 'MTP Registration',
        'pndt_registration': 'PNDT Registration',
        'organ_transplant_registration': 'Organ Transplant Registration',
        'state_clinical_license': 'State Clinical License',
        'nabh_accreditation': 'NABH Accreditation',
        'nabl_accreditation': 'NABL Accreditation',
        'jci_accreditation': 'JCI Accreditation',
        'iso_certification': 'ISO Certification',
        'bio_medical_waste_mgmt_registration': 'Biomedical Waste Management Registration',
        'fire_safety_registration': 'Fire Safety Registration',
        'other_certifications': 'Other Certifications',
        'gst_number': 'GST Number',
        'tan_number': 'TAN Number',
        'pan_number': 'PAN Number',
        'pan_name': 'PAN Name',
        'establishment_year': 'Establishment Year',
        'certificate_of_incorporation': 'Certificate of Incorporation',
        'document_consistency_verified': 'Document Consistency Verified',
        'edited_certificates_found': 'Edited Certificates Found',
        'govt_schemes_run': 'Government Schemes Run',
        'govt_schemes_run_other': 'Other Government Schemes',
        'tpas_empaneled': 'TPAs Empaneled',
        'tpas_empaneled_other': 'Other TPAs Empaneled',
        'insurance_companies_empaneled': 'Insurance Companies Empaneled',
        'corporates_empaneled': 'Corporates Empaneled',
        'medical_superintendent_name': 'Medical Superintendent Name',
        'medical_superintendent_working_type': 'Medical Superintendent Working Type',
        'medical_superintendent_contact': 'Medical Superintendent Contact',
        'medical_director_name': 'Medical Director Name',
        'medical_director_working_type': 'Medical Director Working Type',
        'medical_director_contact': 'Medical Director Contact',
        'hospital_manager_name': 'Hospital Manager Name',
        'hospital_manager_working_type': 'Hospital Manager Working Type',
        'hospital_manager_contact': 'Hospital Manager Contact',
        'tpa_insurance_incharge_name': 'TPA/Insurance Incharge Name',
        'tpa_insurance_incharge_working_type': 'TPA/Insurance Incharge Working Type',
        'tpa_insurance_incharge_contact': 'TPA/Insurance Incharge Contact',
        'nursing_incharge_name': 'Nursing Incharge Name',
        'nursing_incharge_working_type': 'Nursing Incharge Working Type',
        'nursing_incharge_contact': 'Nursing Incharge Contact',
        'surgeon_name': 'Surgeon Name',
        'surgeon_working_type': 'Surgeon Working Type',
        'surgeon_contact': 'Surgeon Contact',
        'physician_name': 'Physician Name',
        'physician_working_type': 'Physician Working Type',
        'physician_contact': 'Physician Contact',
        'intensivist_name': 'Intensivist Name',
        'intensivist_working_type': 'Intensivist Working Type',
        'intensivist_contact': 'Intensivist Contact',
        'pathologist_name': 'Pathologist Name',
        'pathologist_working_type': 'Pathologist Working Type',
        'pathologist_contact': 'Pathologist Contact',
        'lab_technician_name': 'Lab Technician Name',
        'lab_technician_working_type': 'Lab Technician Working Type',
        'lab_technician_contact': 'Lab Technician Contact',
        'radiologist_name': 'Radiologist Name',
        'radiologist_working_type': 'Radiologist Working Type',
        'radiologist_contact': 'Radiologist Contact',
        'key_doctor_registration_numbers': 'Key Doctor Registration Numbers',
        'clinician_sample_registrations': 'Clinician Sample Registrations',
        'nurse_sample_registrations': 'Nurse Sample Registrations',
        'number_of_specialist_doctors': 'Number of Specialist Doctors',
        'number_of_super_specialist_doctors': 'Number of Super Specialist Doctors',
        'number_of_mbbs_and_above_doctors': 'Number of MBBS and Above Doctors',
        'number_of_trained_nurses': 'Number of Trained Nurses',
        'number_of_housemen': 'Number of Housemen',
        'number_of_resident_doctors': 'Number of Resident Doctors',
        'number_of_untrained_nurses': 'Number of Untrained Nurses',
        'number_of_admin_staff': 'Number of Admin Staff',
        'number_of_other_staff': 'Number of Other Staff',
        'total_number_of_staff': 'Total Number of Staff',
        'nurse_to_patient_ratio': 'Nurse to Patient Ratio',
        'type_of_building': 'Type of Building',
        'area_in_square_feet': 'Area (Sq. Ft.)',
        'number_of_floors': 'Number of Floors',
        'number_of_beds': 'Number of Beds',
        'licensed_bed_count': 'Licensed Bed Count',
        'in_house_lab': 'In-House Lab',
        'in_house_pharmacy': 'In-House Pharmacy',
        'ambulance_service': 'Ambulance Service',
        'reception': 'Reception',

        'tpa_insurance_department': 'TPA/Insurance Department',
        'patient_counselling_area': 'Patient Counselling Area',
        'physiotherapy_department': 'Physiotherapy Department',
        'admin_office_present': 'Admin Office Present',
        'number_of_opds': 'Number of OPDs',
        'number_of_icus': 'Number of ICUs',
        'total_icu_beds': 'Total ICU Beds',
        'special_rooms': 'Special Rooms',
        'semi_special_rooms': 'Semi Special Rooms',
        'number_of_suites': 'Number of Suites',
        'number_of_general_wards': 'Number of General Wards',
        'number_of_general_ward_beds': 'General Ward Bed Count',
        'types_of_operation_theaters': 'Types of Operation Theaters',
        'number_of_operation_theaters': 'Number of Operation Theaters',
        'central_oxygen_supply': 'Central Oxygen Supply',
        'number_of_ventilators': 'Number of Ventilators',
        'infra_equipment_validation': 'Infrastructure Equipment Validation',
        'infrastructure_verified': 'Infrastructure Verified',
        'other_infrastructure': 'Other Infrastructure',
        'lab_name': 'Lab Name',
        'lab_pathologist_name': 'Lab Pathologist Name',
        'lab_pathologist_type': 'Lab Pathologist Type',
        'qualification_of_lab_pathologist': 'Qualification of Lab Pathologist',
        'pathologist_own_lab_name_location': 'Own Lab Name and Location',
        'lab_equipment_available': 'Lab Equipment Available',
        'type_of_tests_done': 'Types of Tests Done',
        'tests_outsourced': 'Tests Outsourced',
        'xray_machine': 'X-Ray Machine',
        'c_arm_machine': 'C-Arm Machine',
        'ct_scan': 'CT Scan',
        'mri_machine': 'MRI Machine',
        'usg_machine': 'USG Machine',
        'other_diagnostic_equipment': 'Other Diagnostic Equipment',
        'pharmacy_name': 'Pharmacy Name',
        'pharmacy_drug_license_number': 'Pharmacy Drug License Number',
        'pharmacy_working_hours': 'Pharmacy Working Hours',
        'pharmacist_name': 'Pharmacist Name',
        'pharmacist_registration_number': 'Pharmacist Registration Number',
        'stock_register': 'Stock Register Maintained',
        'purchase_bills': 'Purchase Bills Available',
        'sales_register': 'Sales Register Maintained',
        'other_pharmacy_documentation': 'Other Pharmacy Documentation',
        'ipd_register': 'IPD Register Maintained',
        'icu_register': 'ICU Register Maintained',
        'opd_register': 'OPD Register Maintained',
        'operation_theater_register': 'OT Register Maintained',
        'lab_register': 'Lab Register Maintained',
        'mlc_register': 'MLC Register Maintained',
        'infection_control_protocols': 'Infection Control Protocols Followed',
        'consultant_visits_register': 'Consultant Visits Register Maintained',
        'salary_slips_available': 'Salary Slips Available',
        'salary_payment_mode': 'Salary Payment Mode',

        'bank_account_number': 'Bank Account Number',
        'ifsc_code': 'IFSC Code',
        'cancelled_cheque_copy': 'Cancelled Cheque Copy',

        # BAGIC-specific fields
        'bagic_to_other_patients_ratio': 'BAGIC to Other Patients Ratio',
        'average_stay': 'Average Stay (in Days)',
        'surgical_to_medical_packages_ratio': 'Surgical to Medical Packages Ratio',
        'bagic_mou_copy': 'BAGIC MOU Copy',
        'bagic_soc_type': 'BAGIC SOC Type',
        'bagic_empanelment_done_by': 'BAGIC Empanelment Done By',
        'bagic_empanelment_experience': 'BAGIC Empanelment Experience',
        'bagic_claim_approval_experience': 'BAGIC Claim Approval Experience',
        'bagic_claim_verification_experience': 'BAGIC Claim Verification Experience',
        'bagic_payments_experience': 'BAGIC Payments Experience',
        'suggestion_for_bagic': 'Suggestions for BAGIC',

        # Visit details
        'date_of_visit': 'Date of Visit',
        'visit_done_by': 'Visit Done By',
        'attended_by': 'Attended By',
        'designation': 'Designation',
        'information_given_by': 'Information Given By',
        'remark_by_auditor': 'Remark by Auditor',
        'remark_by_hospital_manager_owner': 'Remark by Hospital Manager/Owner',
        'hospital_manager_owner_name': 'Hospital Manager/Owner Name',

        # Verification flags
        'new_setup_claim_verified': 'New Setup Claim Verified',
        'human_presence_verified': 'Human Presence Verified',
        'local_market_intel_remarks': 'Local Market Intelligence Remarks',
        'repeat_tpa_fraud_suspected': 'Repeat TPA Fraud Suspected',
    }



def get_excluded_fields():
    """Fields to exclude from the report"""
    return {'hospital_id', 'audit_status'}


def format_field_value(value):
    """Format field values for professional display"""
    if value is None or value == "":
        return "Not Specified"
    if isinstance(value, str):
        # Convert abbreviations to full forms
        if value.lower() == 'per':
            return "Private"
        elif value.lower() == 'gov':
            return "Government"
        elif value.lower() == 'city':
            return "Urban"
        elif value.lower() == 'rural':
            return "Rural"
    return str(value).title() if isinstance(value, str) else str(value)


def create_docx_from_hospital(hospital: FullHospitalData) -> BytesIO:
    """Create a professional DOCX report for a hospital"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Main Title
    title = doc.add_heading("HOSPITAL REPORT", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Hospital Name as subtitle
    if hospital.basic_info and hospital.basic_info.hospital_name:
        subtitle = doc.add_heading(hospital.basic_info.hospital_name.upper(), level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add report generation date
    date_para = doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add a line break
    doc.add_paragraph()

    if hospital.basic_info:
        # Create professional field mapping
        field_mapping = create_professional_field_mapping()
        excluded_fields = get_excluded_fields()

        # Basic Information Section
        doc.add_heading("", level=2)

        # Create a table for better formatting
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'Table Grid'

        # Set table column widths
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(4)

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Information'

        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add hospital data to table
        hospital_dict = hospital.basic_info.dict()

        for key, value in hospital_dict.items():
            # Skip excluded fields
            if key in excluded_fields:
                continue

            # Skip None or empty values for cleaner report
            if value is None or value == "":
                continue

            # Add row to table
            row_cells = table.add_row().cells
            row_cells[0].text = field_mapping.get(key, key.replace('_', ' ').title())
            row_cells[1].text = format_field_value(value)

    # Add additional sections if more data is available
    if hasattr(hospital, 'services') and hospital.services:
        doc.add_heading("SERVICES", level=2)
        for service in hospital.services:
            doc.add_paragraph(f"\u2022 {service}", style='List Bullet')

    if hasattr(hospital, 'facilities') and hospital.facilities:
        doc.add_heading("FACILITIES", level=2)
        for facility in hospital.facilities:
            doc.add_paragraph(f"\u2022 {facility}", style='List Bullet')


    # Add footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "This report is generated automatically. For any queries, please contact the administration.")
    footer_run.font.size = Pt(10)
    footer_run.italic = True

    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def export_hospitals_to_docx_zip(hospitals: List[FullHospitalData]):
    """Export multiple hospitals to a professional DOCX zip file"""
    zip_stream = BytesIO()

    with ZipFile(zip_stream, "w") as zip_file:
        for idx, hospital in enumerate(hospitals):
            doc_io = create_docx_from_hospital(hospital)

            # Create professional filename
            if hospital.basic_info and hospital.basic_info.hospital_name:
                # Clean filename by removing special characters
                hospital_name = hospital.basic_info.hospital_name.replace(" ", "_")
                hospital_name = "".join(c for c in hospital_name if c.isalnum() or c in "._-")
                filename = f"Hospital_Report_{hospital_name}.docx"
            else:
                filename = f"Hospital_Report_{idx + 1:03d}.docx"

            zip_file.writestr(filename, doc_io.read())

    zip_stream.seek(0)
    return StreamingResponse(
        zip_stream,
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename=Hospital_Reports_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        }
    )
